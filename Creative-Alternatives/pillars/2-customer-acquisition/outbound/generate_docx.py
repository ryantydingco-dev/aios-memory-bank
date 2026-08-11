"""Generate .docx for Creative Alternatives Youth Sports Clubs campaign."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Creative Alternatives — Youth Sports Clubs Campaign V1', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = [
    ("Status", "Ready to build lists + load into SmartLead"),
    ("Created", "2026-04-16"),
    ("Senders", "Split test — Maclaine (50%) / Ryan (50%)"),
    ("Target", "Youth sports clubs, athletic academies, travel teams, boutique fitness"),
    ("Offer", "Free custom branded online store — full fulfillment, club earns on every order"),
    ("Proof Point", "Farm & Forge Club (active client, College Grove TN)"),
]
for label, value in meta:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

# ══════════════════════════════════════════
# WHY THIS VERTICAL
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Why Youth Sports Clubs', level=1)

doc.add_paragraph(
    'The summer camps campaign hit a 10.1% reply rate — the best-performing campaign across '
    'all of Creative Alternatives\' outreach. Youth sports clubs share the same dynamics that '
    'made camps work:'
)

reasons = [
    ("Same buyer psychology.", "Parents spending on their kids, community identity, pride in the organization."),
    ("Built-in demand.", "Families, teams, tournaments, and seasonal registrations all create natural buying moments for branded gear."),
    ("Same zero-risk offer.", "Free store, no inventory, earn on every order. The offer sells itself."),
    ("Massive market.", "Tens of thousands of clubs across the US — tennis, swim, gymnastics, martial arts, dance, travel ball, CrossFit, and more."),
    ("Proof point ready.", "Farm & Forge Club is already an active client. Real orders, real store, real results to reference."),
]
for bold_part, rest in reasons:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part + " ")
    run.bold = True
    p.add_run(rest)

# ══════════════════════════════════════════
# WHO WE'RE TARGETING
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Who We\'re Targeting', level=1)

doc.add_paragraph('We\'re going broad across all youth sports and athletic club types:')

club_types = [
    "Tennis clubs & academies",
    "Swim clubs & aquatics programs",
    "Gymnastics studios",
    "Martial arts / BJJ / karate / taekwondo",
    "Dance studios & academies",
    "Travel/club baseball, soccer, lacrosse (AAU-style)",
    "CrossFit boxes & boutique fitness studios",
    "Volleyball, hockey, wrestling, fencing clubs",
    "Cheer gyms",
    "Track, running, cycling, triathlon clubs",
    "Youth basketball & soccer academies",
]
for club in club_types:
    doc.add_paragraph(club, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Who We\'re Contacting', level=2)
doc.add_paragraph(
    'Decision makers who control merch and marketing — owners, directors, head coaches, '
    'GMs, marketing managers. Not assistant coaches or front desk staff.'
)

doc.add_heading('Buying Signals We\'re Looking For', level=2)
signals = [
    ("No online store", "Club has no merch store on their website — they don't have this yet, and we're offering it for free."),
    ("Already sells merch", "Club has a pro shop or spirit wear — they're already in the habit, we're offering a better/easier way."),
    ("Large membership", "Multiple programs, teams, or locations — higher volume potential."),
    ("Hosts events", "Tournaments, showcases, meets, recitals — event-specific merch opportunity."),
]
for sig, desc in signals:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(sig + " — ")
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════
# SPLIT TEST
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Split Test Setup', level=1)

doc.add_paragraph(
    'We\'re running a 50/50 split test between two senders to find out which voice performs '
    'better with this audience. Both use the same structure and angles — the difference is tone.'
)

split_table = doc.add_table(rows=3, cols=3)
split_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['', 'Variant A — Maclaine', 'Variant B — Ryan']):
    split_table.rows[0].cells[i].text = h

split_data = [
    ('Tone', 'Warm, casual, personal — "do you guys sell gear?"', 'Professional but friendly — "does {{company_name}} currently sell merchandise?"'),
    ('Best for', 'Smaller clubs, family-run, community feel', 'Larger academies, multi-location, more corporate'),
]
for row_idx, (label, a, b) in enumerate(split_data, 1):
    split_table.rows[row_idx].cells[0].text = label
    split_table.rows[row_idx].cells[1].text = a
    split_table.rows[row_idx].cells[2].text = b

doc.add_paragraph()
doc.add_paragraph(
    'Subject lines are identical across both variants so we can isolate sender performance. '
    'After 2-3 weeks we\'ll have enough data to know which sender wins and can consolidate.'
)

# ══════════════════════════════════════════
# SEQUENCE OVERVIEW
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Email Sequence Overview', level=1)

seq_table = doc.add_table(rows=6, cols=3)
seq_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Email', 'Purpose', 'Delay']):
    seq_table.rows[0].cells[i].text = h

emails_overview = [
    ('1', 'The Opener — do you sell gear? (curiosity, low friction)', 'Day 0'),
    ('2', 'The Proof — Farm & Forge case study', 'Day 3'),
    ('3', 'The Fundraising Angle — passive revenue for the club', 'Day 7'),
    ('4', 'The Seasonal Hook — get ready before next season', 'Day 12'),
    ('5', 'The Breakup — friendly close, leave door open', 'Day 18'),
]
for row_idx, (num, purpose, delay) in enumerate(emails_overview, 1):
    seq_table.rows[row_idx].cells[0].text = num
    seq_table.rows[row_idx].cells[1].text = purpose
    seq_table.rows[row_idx].cells[2].text = delay

# ══════════════════════════════════════════
# EMAIL COPY
# ══════════════════════════════════════════

def add_email_split(doc, number, name, subjects, body_a, body_b):
    doc.add_page_break()
    doc.add_heading(f'Email {number}: {name}', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Subject line options (test 2-3):')
    run.bold = True
    for s in subjects:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('Variant A (Maclaine)', level=2)
    for line in body_a.strip().split('\n'):
        doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_heading('Variant B (Ryan)', level=2)
    for line in body_b.strip().split('\n'):
        doc.add_paragraph(line)


add_email_split(doc, 1, 'The Opener',
    ['free store for {{company_name}}', 'quick question about {{company_name}} gear', 'does {{company_name}} sell merch?'],
    """Hi {{first_name}},

Do you guys currently sell any branded gear to your members and families?

We build free custom online stores for sports clubs — we handle all the design, fulfillment, and shipping, and your club earns a cut on every order. No inventory, no upfront cost, nothing to manage.

Just curious if it's something you've thought about.

Maclaine""",
    """Hi {{first_name}},

Does {{company_name}} currently sell any branded merchandise to members and families?

We set up free custom online stores for sports clubs and academies — we handle design, fulfillment, and shipping. Your club earns on every order with zero inventory or upfront cost.

Wanted to see if this is something that's been on your radar.

Ryan""")


add_email_split(doc, 2, 'The Proof (Farm & Forge)',
    ['re: free store for {{company_name}}', 'what we just built for a club in Tennessee', 'quick example'],
    """Hi {{first_name}},

Wanted to share a quick example — we just built a branded store for Farm & Forge Club, a youth athletics academy in Tennessee. Their families can now order custom gear year-round, and the club earns on every sale without lifting a finger.

We handle the whole thing — design, printing, fulfillment, shipping. They just shared the link with their community and orders started coming in.

Would love to do something similar for {{company_name}} if you're interested. Happy to show you what their store looks like.

Maclaine""",
    """Hi {{first_name}},

Quick example of what this looks like in practice — we recently set up a branded online store for Farm & Forge Club, a youth sports academy in Tennessee.

Their members and families can browse and order custom gear anytime. The club earns a percentage on every order. We handle design, printing, fulfillment, and shipping — they just promote the link.

If {{company_name}} would benefit from something similar, I'm happy to walk you through how it works. Can also send over what their store looks like.

Ryan""")


add_email_split(doc, 3, 'The Fundraising Angle',
    ['{{company_name}} merch as a fundraiser', 'passive revenue idea for {{company_name}}', 'your members already want this'],
    """Hi {{first_name}},

One thing a lot of clubs don't realize — a branded merch store can quietly become one of your best fundraisers.

Parents and families love buying gear with their club's logo. They're going to buy t-shirts and hoodies anyway — might as well be yours. And when every order puts money back into {{company_name}} with zero work on your end, it's kind of a no-brainer.

We handle everything — you just share the store link. Want me to put together a quick mockup of what a {{company_name}} store could look like?

Maclaine""",
    """Hi {{first_name}},

Something worth considering — a branded merch store can double as a passive fundraiser for {{company_name}}.

Members and families are already buying athletic gear. When it's branded with your club's logo and every order generates revenue back to you — with zero inventory, zero fulfillment, and zero upfront cost — it becomes one of the easiest revenue streams you can add.

Happy to put together a quick mockup of what a {{company_name}} store would look like. Just reply and I'll send it over.

Ryan""")


add_email_split(doc, 4, 'The Seasonal Hook',
    ['getting ahead of next season', '{{company_name}} gear for the fall', 'quick thought on merch timing'],
    """Hi {{first_name}},

Most clubs we work with launch their store right before a new season starts — it's when families are most excited and most likely to buy.

If you've been thinking about merch at all, now is a great time to get set up so you're ready when the next season kicks off. We can have a full store designed and live in about a week.

Just curious if the timing makes sense for you guys?

Maclaine""",
    """Hi {{first_name}},

Quick thought — the clubs that get the most out of their merch stores are the ones that launch right before a new season or registration period.

If {{company_name}} has a season coming up, we can have a full branded store designed and live within a week. Families are most likely to buy when they're already excited about the program.

Would it make sense to get something set up ahead of your next season?

Ryan""")


add_email_split(doc, 5, 'The Breakup',
    ['last note from me', 'closing the loop', 'no worries at all'],
    """Hi {{first_name}},

Totally understand if this isn't a priority right now — just wanted to make sure the offer was on your radar.

If you ever want to explore a branded store for {{company_name}}, we'd love to help. It's completely free to set up and there's never any obligation.

Wishing you a great season!

Maclaine""",
    """Hi {{first_name}},

I'll assume the timing isn't right — completely understand.

If {{company_name}} ever wants to explore a branded merch store, the offer stands. Free to set up, no obligation, and we handle everything.

Wishing you and your team a great season.

Ryan""")

# ══════════════════════════════════════════
# TONE NOTES
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Tone & Strategy Notes', level=1)

notes = [
    ("Keep it casual.", "These are club owners and coaches, not corporate executives. The summer camps campaign worked because it was short, direct, and zero-pressure."),
    ("The offer sells itself.", "Free store, no inventory, earn on every order. Don't over-explain — just make sure they understand there's no risk."),
    ("Farm & Forge is real proof.", "Email 2 references a real client with real orders. This is what separates us from a generic pitch."),
    ("Fundraising > promotional products.", "\"Your club earns on every order\" is more compelling than \"we sell promotional products.\" Frame it as revenue, not merch."),
    ("Seasonal urgency is natural.", "Email 4 ties to the rhythm of sports seasons — this creates real timing pressure without being salesy."),
    ("Split test will reveal the winner.", "After 2-3 weeks, consolidate to whichever sender performs better. Subject lines stay the same across both variants."),
]
for bold_part, rest in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part + " ")
    run.bold = True
    p.add_run(rest)

# ══════════════════════════════════════════
# NEXT STEPS
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_heading('Next Steps', level=1)

steps = [
    "Pull Apollo lists — youth sports clubs, athletic academies, fitness studios across the US",
    "Validate emails via ZeroBounce",
    "Split list 50/50 for Maclaine vs Ryan sender variants",
    "Load into SmartLead as two campaigns (Variant A + Variant B) with same subject lines",
    "Start sending at 50/day per variant (100/day total), ramp after warm-up",
    "Review split test results after 2-3 weeks — consolidate to winning sender",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

# --- Save ---
output_path = '/Users/ryantydingco/Documents/AIOS/aios-starter-kit/clients/creative-alternatives/outbound/Creative Alternatives - Youth Sports Clubs Campaign V1.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
